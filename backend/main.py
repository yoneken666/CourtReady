import uvicorn
from fastapi import FastAPI, UploadFile, File, Depends, HTTPException, status, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import OAuth2PasswordRequestForm
from sqlalchemy.orm import Session
from typing import List, Optional
import os
import shutil
import pypdf
import docx
import re
from dotenv import load_dotenv
from starlette.concurrency import run_in_threadpool

load_dotenv()

import models, schemas, auth
from database import engine, get_db
from caseanalyzer import legal_engine
from casematching import find_similar_cases          # ← NEW
from argumentbuilder import generate_arguments

models.Base.metadata.create_all(bind=engine)

app = FastAPI(debug=True)

origins = ["http://localhost:5173"]
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIRECTORY = "./uploaded_files"
os.makedirs(UPLOAD_DIRECTORY, exist_ok=True)


# --- Helper Functions ---
def clean_extracted_text(text: str) -> str:
    """Removes null bytes and excessive whitespace to prevent AI errors."""
    if not text:
        return ""
    text = text.replace("\x00", "")
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def extract_text_from_file_obj(file_obj, filename: str) -> str:
    text = ""
    try:
        file_obj.seek(0)

        if filename.endswith(".pdf"):
            reader = pypdf.PdfReader(file_obj)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + " "
        elif filename.endswith(".docx"):
            doc = docx.Document(file_obj)
            text = "\n".join([para.text for para in doc.paragraphs])

        return clean_extracted_text(text)
    except Exception as e:
        print(f"Error extracting text from {filename}: {e}")
        return ""


def extract_text_from_file(file_path: str, content_type: str) -> str:
    with open(file_path, "rb") as f:
        return extract_text_from_file_obj(f, file_path)


# --- AUTH ENDPOINTS ---
@app.post("/api/signup", status_code=status.HTTP_201_CREATED)
def create_user(user: schemas.UserCreate, db: Session = Depends(get_db)):
    db_user = db.query(models.User).filter(models.User.email == user.email).first()
    if db_user:
        raise HTTPException(status_code=400, detail="Email already registered")

    hashed_password = auth.get_password_hash(user.password)
    new_user = models.User(email=user.email, hashed_password=hashed_password)
    db.add(new_user)
    db.commit()
    db.refresh(new_user)
    return {"message": "User created successfully"}


@app.post("/api/token", response_model=schemas.Token)
def login_for_access_token(form_data: OAuth2PasswordRequestForm = Depends(), db: Session = Depends(get_db)):
    user = db.query(models.User).filter(models.User.email == form_data.username).first()
    if not user or not auth.verify_password(form_data.password, user.hashed_password):
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Incorrect email or password")
    access_token = auth.create_access_token(data={"sub": user.email})
    return {"access_token": access_token, "token_type": "bearer"}


# --- AI ANALYSIS ENDPOINT ---
@app.post("/api/analyze-case", response_model=schemas.AnalysisResponse)
async def analyze_case(
        caseTitle: str = Form(...),
        caseType: str = Form(...),
        caseDescription: str = Form(...),
        file: Optional[UploadFile] = File(None),
        current_user: dict = Depends(auth.get_current_user)
):
    ALLOWED_CATEGORIES = ["Contract Disputes", "Property Disputes", "Family Disputes"]
    if caseType not in ALLOWED_CATEGORIES:
        raise HTTPException(status_code=400, detail="Invalid Case Type.")

    full_description = caseDescription

    if file:
        print(f"Processing attached file: {file.filename}")
        extracted_text = extract_text_from_file_obj(file.file, file.filename)

        if extracted_text and len(extracted_text) > 10:
            limit = 50000
            if len(extracted_text) > limit:
                extracted_text = extracted_text[:limit] + "... [Text Truncated]"

            full_description += f"\n\n[Attached Document Content]:\n{extracted_text}"
        else:
            print("Warning: Extracted text was empty or too short.")

    try:
        result = await run_in_threadpool(
            legal_engine.analyze_case,
            description=full_description,
            category=caseType
        )
        return schemas.AnalysisResponse(**result)
    except Exception as e:
        print(f"Analysis Error: {e}")
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")


# --- CASE MATCHING ENDPOINT ---                                               ← NEW
@app.post("/api/match-cases", response_model=schemas.CaseMatchResponse)
async def match_cases(
        caseTitle: str = Form(...),
        caseDescription: str = Form(...),
        file: Optional[UploadFile] = File(None),
        current_user: dict = Depends(auth.get_current_user)
):
    """
    Accepts a family dispute description (and optional supporting document),
    finds the top-5 most keyword-similar past cases from ./family_cases/,
    then uses Gemini to contextually rank and filter them.
    Returns up to 3 genuinely similar cases with relevance labels and explanations.
    """
    full_description = caseDescription

    if file:
        print(f"[match-cases] Processing attached file: {file.filename}")
        extracted = extract_text_from_file_obj(file.file, file.filename)
        if extracted and len(extracted) > 10:
            full_description += f"\n\n[Attached Document]:\n{extracted[:30000]}"

    result = await run_in_threadpool(find_similar_cases, full_description)
    return result


# --- CASE MANAGEMENT ENDPOINTS ---
@app.post("/api/case", response_model=schemas.CaseResponse)
def create_case(
        case_details: schemas.CaseIntake,
        current_user: dict = Depends(auth.get_current_user),
        db: Session = Depends(get_db)
):
    user = db.query(models.User).filter(models.User.email == current_user['email']).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    new_case = models.Case(
        title=case_details.caseTitle,
        case_type=case_details.caseType,
        description=case_details.caseDescription,
        user_id=user.id
    )
    db.add(new_case)
    db.commit()
    db.refresh(new_case)

    return {
        "caseTitle": new_case.title,
        "caseType": new_case.case_type,
        "caseDescription": new_case.description,
        "id": new_case.id,
        "owner_email": user.email
    }


@app.post("/api/upload-documents")
def upload_documents(
        case_id: int = Form(...),
        files: List[UploadFile] = File(...),
        current_user: dict = Depends(auth.get_current_user),
        db: Session = Depends(get_db)
):
    user = db.query(models.User).filter(models.User.email == current_user['email']).first()
    case = db.query(models.Case).filter(models.Case.id == case_id, models.Case.user_id == user.id).first()

    if not case:
        raise HTTPException(status_code=404, detail="Case not found or unauthorized")

    saved_files = []

    for file in files:
        file_path = os.path.join(UPLOAD_DIRECTORY, f"case_{case_id}_{file.filename}")
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        extracted_text = extract_text_from_file(file_path, file.content_type)

        new_doc = models.Document(
            filename=file.filename,
            file_path=file_path,
            extracted_text=extracted_text,
            case_id=case.id
        )
        db.add(new_doc)
        saved_files.append({"filename": file.filename})

    db.commit()
    return {"status": "success", "saved_files": saved_files}

@app.post("/api/build-arguments")
async def build_arguments(
        caseDescription: str = Form(...),
        analysisResult:  str = Form(...),
        matchingResult:  str = Form(...),
        current_user: dict = Depends(auth.get_current_user)
):
    import json as _json
    try:
        analysis = _json.loads(analysisResult)
        matching = _json.loads(matchingResult)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid JSON in analysisResult or matchingResult.")
    try:
        result = await run_in_threadpool(
            generate_arguments,
            case_description=caseDescription,
            analysis_result=analysis,
            matching_result=matching,
        )
        return result
    except Exception as e:
        print(f"Argument Builder Error: {e}")
        raise HTTPException(status_code=500, detail=f"Argument generation failed: {str(e)}")


if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
